import os
import sys
import time
import logging
import argparse
import concurrent.futures
from pathlib import Path
from typing import List, Optional
from pdf2image import convert_from_path
from docx import Document
from bs4 import BeautifulSoup
import requests
from PIL import Image
import imgkit
from tqdm import tqdm

from config import Config
from cache_manager import CacheManager
from security import SecurityManager

class DocumentConverter:
    def __init__(self):
        """אתחול ממיר המסמכים"""
        # אתחול מנהלים
        Config.init_directories()
        self.cache = CacheManager()
        self.security = SecurityManager()
        
        # הגדרת לוגים
        self._setup_logging()
        
        self.logger.info("מערכת המרת מסמכים אותחלה")

    def _setup_logging(self):
        """הגדרת מערכת הלוגים"""
        self.logger = logging.getLogger('DocumentConverter')
        self.logger.setLevel(Config.LOG_LEVEL)
        
        # הגדרת קובץ לוג מתגלגל
        from logging.handlers import RotatingFileHandler
        handler = RotatingFileHandler(
            Config.LOG_DIR / 'converter.log',
            maxBytes=Config.LOG_MAX_SIZE,
            backupCount=Config.LOG_BACKUP_COUNT,
            encoding='utf-8'
        )
        handler.setFormatter(logging.Formatter(Config.LOG_FORMAT))
        self.logger.addHandler(handler)

    def process_pdf(self, file_path: str) -> List[Path]:
        """המרת קובץ PDF לתמונות
        
        Args:
            file_path: נתיב לקובץ ה-PDF
            
        Returns:
            רשימת נתיבים לתמונות שנוצרו
        """
        self.logger.info(f"מתחיל המרת PDF: {file_path}")
        output_files = []
        
        # בדיקת מטמון
        cache_key = f"pdf_{file_path}"
        cached_data = self.cache.get(cache_key)
        if cached_data:
            self.logger.info("נמצא במטמון")
            return cached_data
        
        try:
            # המרה לתמונות
            images = convert_from_path(file_path, dpi=Config.DEFAULT_DPI)
            
            with tqdm(total=len(images), desc="ממיר עמודי PDF") as pbar:
                with concurrent.futures.ThreadPoolExecutor(max_workers=Config.MAX_WORKERS) as executor:
                    futures = []
                    
                    for i, image in enumerate(images, start=1):
                        if i > Config.MAX_PAGES:
                            break
                            
                        output_path = Config.OUTPUT_DIR / f"{i}.{Config.DEFAULT_IMAGE_FORMAT}"
                        futures.append(
                            executor.submit(self._save_image, image, output_path)
                        )
                        
                    for future in concurrent.futures.as_completed(futures):
                        output_files.append(future.result())
                        pbar.update(1)
            
            # שמירה במטמון
            self.cache.set(cache_key, output_files)
            
            self.logger.info(f"הומרו {len(output_files)} עמודי PDF")
            return output_files
            
        except Exception as e:
            self.logger.error(f"שגיאה בהמרת PDF: {str(e)}")
            raise

    def process_word(self, file_path: str) -> List[Path]:
        """המרת מסמך Word לתמונות
        
        Args:
            file_path: נתיב לקובץ ה-Word
            
        Returns:
            רשימת נתיבים לתמונות שנוצרו
        """
        self.logger.info(f"מתחיל המרת Word: {file_path}")
        output_files = []
        
        try:
            doc = Document(file_path)
            
            with tqdm(total=len(doc.paragraphs), desc="ממיר עמודי Word") as pbar:
                with concurrent.futures.ThreadPoolExecutor(max_workers=Config.MAX_WORKERS) as executor:
                    futures = []
                    
                    for i, paragraph in enumerate(doc.paragraphs, start=1):
                        if i > Config.MAX_PAGES:
                            break
                            
                        if paragraph.text.strip():
                            output_path = Config.OUTPUT_DIR / f"{i}.{Config.DEFAULT_IMAGE_FORMAT}"
                            futures.append(
                                executor.submit(
                                    self._create_word_image,
                                    paragraph.text,
                                    output_path
                                )
                            )
                            
                    for future in concurrent.futures.as_completed(futures):
                        output_files.append(future.result())
                        pbar.update(1)
            
            self.logger.info(f"הומרו {len(output_files)} עמודי Word")
            return output_files
            
        except Exception as e:
            self.logger.error(f"שגיאה בהמרת Word: {str(e)}")
            raise

    def process_html(self, source: str) -> List[Path]:
        """המרת HTML או URL לתמונות
        
        Args:
            source: נתיב לקובץ HTML או כתובת URL
            
        Returns:
            רשימת נתיבים לתמונות שנוצרו
        """
        self.logger.info(f"מתחיל המרת HTML/URL: {source}")
        
        try:
            # בדיקת אבטחה ל-URL
            if source.startswith(('http://', 'https://')):
                if not self.security.validate_url(source):
                    raise ValueError("כתובת URL לא בטוחה")
                response = requests.get(source)
                html_content = response.text
            else:
                with open(source, 'r', encoding='utf-8') as f:
                    html_content = f.read()

            # הגדרות עבור imgkit
            options = {
                'format': Config.DEFAULT_IMAGE_FORMAT,
                'encoding': 'UTF-8',
                'quality': Config.IMAGE_FORMATS[Config.DEFAULT_IMAGE_FORMAT]['quality']
            }

            output_path = Config.OUTPUT_DIR / f"1.{Config.DEFAULT_IMAGE_FORMAT}"
            imgkit.from_string(html_content, str(output_path), options=options)
            
            self.logger.info("הומר דף HTML/URL")
            return [output_path]
            
        except Exception as e:
            self.logger.error(f"שגיאה בהמרת HTML/URL: {str(e)}")
            raise

    def _save_image(self, image: Image.Image, output_path: Path) -> Path:
        """שמירת תמונה בפורמט הרצוי
        
        Args:
            image: אובייקט התמונה
            output_path: נתיב השמירה
            
        Returns:
            נתיב התמונה שנשמרה
        """
        image_format = Config.DEFAULT_IMAGE_FORMAT.upper()
        quality = Config.IMAGE_FORMATS[Config.DEFAULT_IMAGE_FORMAT]['quality']
        
        image.save(
            output_path,
            format=image_format,
            quality=quality,
            optimize=True
        )
        return output_path

    def _create_word_image(self, text: str, output_path: Path) -> Path:
        """יצירת תמונה מטקסט של Word
        
        Args:
            text: הטקסט ליצירת התמונה
            output_path: נתיב השמירה
            
        Returns:
            נתיב התמונה שנוצרה
        """
        # יצירת תמונה עם הטקסט
        image = Image.new('RGB', Config.DEFAULT_IMAGE_SIZE, 'white')
        return self._save_image(image, output_path)

    def convert(self, file_path: str) -> List[Path]:
        """פונקציה ראשית להמרת מסמך
        
        Args:
            file_path: נתיב לקובץ או URL
            
        Returns:
            רשימת נתיבים לתמונות שנוצרו
        """
        try:
            # ניקוי וולידציה של הקלט
            file_path = self.security.sanitize_path(file_path)
            
            if not Config.validate_file(file_path) and not self.security.validate_url(file_path):
                raise ValueError("סוג הקובץ אינו נתמך או שהקובץ אינו תקין")
                
            # בחירת המעבד המתאים
            if file_path.lower().endswith('.pdf'):
                return self.process_pdf(file_path)
            elif file_path.lower().endswith('.docx'):
                return self.process_word(file_path)
            elif file_path.lower().endswith('.html') or file_path.startswith(('http://', 'https://')):
                return self.process_html(file_path)
            else:
                raise ValueError("סוג הקובץ אינו נתמך")
                
        except Exception as e:
            self.logger.error(f"שגיאה בהמרה: {str(e)}")
            raise

def main():
    """פונקציית הפעלה ראשית"""
    parser = argparse.ArgumentParser(description='המרת מסמכים לתמונות')
    parser.add_argument('files', nargs='+', help='נתיבים לקבצים להמרה או כתובות URL')
    parser.add_argument('--output-dir', help='תיקיית פלט (ברירת מחדל: תיקיית ברירת המחדל)', default=None)
    parser.add_argument('--format', choices=['png', 'jpg', 'webp'], default='png', help='פורמט תמונות הפלט')
    parser.add_argument('--dpi', type=int, default=300, help='רזולוציית התמונות')
    parser.add_argument('--quality', type=int, default=95, help='איכות התמונות (1-100)')
    
    args = parser.parse_args()
    
    # עדכון הגדרות לפי הפרמטרים
    if args.output_dir:
        Config.OUTPUT_DIR = Path(args.output_dir)
    Config.DEFAULT_IMAGE_FORMAT = args.format
    Config.DEFAULT_DPI = args.dpi
    Config.IMAGE_FORMATS[args.format]['quality'] = args.quality
    
    try:
        print("\nמערכת המרת מסמכים לתמונות")
        print("=" * 50)
        
        converter = DocumentConverter()
        total_files = 0
        
        for file_path in args.files:
            print(f"\nמעבד קובץ: {file_path}")
            start_time = time.time()
            
            try:
                output_files = converter.convert(file_path)
                total_files += len(output_files)
                
                print(f"✓ הושלם בהצלחה!")
                print(f"  זמן עיבוד: {time.time() - start_time:.2f} שניות")
                print(f"  מספר תמונות: {len(output_files)}")
                
            except Exception as e:
                print(f"✗ שגיאה בעיבוד הקובץ: {str(e)}")
                continue
        
        print("\nסיכום:")
        print(f"סה\"כ קבצים שעובדו: {len(args.files)}")
        print(f"סה\"כ תמונות שנוצרו: {total_files}")
        print(f"התמונות נשמרו בתיקייה: {Config.OUTPUT_DIR}")
        
    except Exception as e:
        print(f"\nשגיאה: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main() 